#!/usr/bin/env python3
"""Regression tests for inspect_book.py using temporary generated books."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
import subprocess
import sys
import tempfile
import unittest
import zipfile


SCRIPT = Path(__file__).with_name("inspect_book.py")


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def write_epub(path: Path) -> None:
    container = """<?xml version="1.0"?>
<container xmlns="urn:oasis:names:tc:opendocument:xmlns:container" version="1.0">
  <rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>"""
    opf = """<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>测试 EPUB</dc:title><dc:creator>测试作者</dc:creator>
  </metadata>
  <manifest>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="c1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>
    <item id="c2" href="chapter2.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine><itemref idref="c1"/><itemref idref="c2"/></spine>
</package>"""
    nav = """<html xmlns="http://www.w3.org/1999/xhtml"><body><nav><ol>
<li><a href="chapter1.xhtml">第一章</a></li><li><a href="chapter2.xhtml">第二章</a></li>
</ol></nav></body></html>"""
    chapter1 = """<html xmlns="http://www.w3.org/1999/xhtml"><head><title>第一章</title></head>
<body><h1>第一章</h1><h2>核心概念</h2><p>这是可读取的正文，用于验证章节顺序与定位。</p></body></html>"""
    chapter2 = """<html xmlns="http://www.w3.org/1999/xhtml"><head><title>第二章</title></head><body></body></html>"""
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        archive.writestr("META-INF/container.xml", container)
        archive.writestr("OEBPS/content.opf", opf)
        archive.writestr("OEBPS/nav.xhtml", nav)
        archive.writestr("OEBPS/chapter1.xhtml", chapter1)
        archive.writestr("OEBPS/chapter2.xhtml", chapter2)


def write_role_epub(path: Path, include_cover: bool = True) -> None:
    container = """<?xml version="1.0"?>
<container xmlns="urn:oasis:names:tc:opendocument:xmlns:container" version="1.0">
  <rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>"""
    opf = """<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:title>角色测试书</dc:title></metadata>
  <manifest>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="cover" href="cover.xhtml" media-type="application/xhtml+xml"/>
    <item id="chapter" href="chapter.xhtml" media-type="application/xhtml+xml"/>
    <item id="license" href="license.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine><itemref idref="cover"/><itemref idref="chapter"/><itemref idref="license"/></spine>
</package>"""
    nav = """<html xmlns="http://www.w3.org/1999/xhtml"><body><nav><ol>
<li><a href="cover.xhtml">Cover</a></li><li><a href="chapter.xhtml">第一章</a></li>
<li><a href="license.xhtml">Project License</a></li></ol></nav></body></html>"""
    cover = """<html xmlns="http://www.w3.org/1999/xhtml"><head><title>Cover</title></head><body></body></html>"""
    chapter = """<html xmlns="http://www.w3.org/1999/xhtml"><body><h1>第一章</h1><p>这是全书核心正文。</p></body></html>"""
    license_text = "许可条款不应被当作作者论证章节。"
    license_page = f"""<html xmlns="http://www.w3.org/1999/xhtml"><body><h1>Project License</h1><p>{license_text}</p></body></html>"""
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        archive.writestr("META-INF/container.xml", container)
        archive.writestr("OEBPS/content.opf", opf)
        archive.writestr("OEBPS/nav.xhtml", nav)
        if include_cover:
            archive.writestr("OEBPS/cover.xhtml", cover)
        archive.writestr("OEBPS/chapter.xhtml", chapter)
        archive.writestr("OEBPS/license.xhtml", license_page)


class InspectBookTests(unittest.TestCase):
    maxDiff = None

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory(prefix="deep-reading-coach-test-")
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def inspect(self, source: Path, name: str) -> tuple[subprocess.CompletedProcess[str], dict, Path]:
        output = self.root / f"out-{name}"
        before = sha256(source)
        process = subprocess.run(
            [sys.executable, str(SCRIPT), str(source), "--output-dir", str(output)],
            text=True,
            capture_output=True,
            check=False,
        )
        self.assertEqual(before, sha256(source), "The source book was modified")
        report_path = output / "inspection.json"
        self.assertTrue(report_path.exists(), process.stderr)
        report = json.loads(report_path.read_text(encoding="utf-8"))
        return process, report, output

    def test_txt_without_headings_uses_line_ranges(self) -> None:
        source = self.root / "plain.txt"
        source.write_text("第一行内容\n第二行内容\n", encoding="utf-8")
        process, report, output = self.inspect(source, "txt")
        self.assertEqual(process.returncode, 0)
        self.assertEqual(report["format"], "txt")
        self.assertEqual(report["directory_status"], "not_applicable")
        self.assertIn("源文本行 1-2", (output / "content.md").read_text(encoding="utf-8"))

    def test_markdown_headings_are_preserved(self) -> None:
        source = self.root / "book.md"
        source.write_text("# 第一部分\n正文甲。\n## 第二节\n正文乙。\n", encoding="utf-8")
        _, report, output = self.inspect(source, "md")
        self.assertEqual(report["status"], "readable")
        self.assertEqual(report["directory_status"], "available")
        self.assertIn("Markdown: 第一部分", (output / "toc.md").read_text(encoding="utf-8"))

    def test_html_title_and_headings(self) -> None:
        source = self.root / "book.html"
        source.write_text(
            "<html><head><title>HTML 书</title><style>hidden</style></head>"
            "<body><h1>开篇</h1><p>正文内容足以读取。</p><script>ignored()</script></body></html>",
            encoding="utf-8",
        )
        _, report, output = self.inspect(source, "html")
        self.assertEqual(report["metadata"]["title"], "HTML 书")
        content = (output / "content.md").read_text(encoding="utf-8")
        self.assertIn("HTML: 开篇", content)
        self.assertNotIn("ignored()", content)
        self.assertNotIn("hidden", content)

    def test_epub_spine_order_and_empty_chapter_warning(self) -> None:
        source = self.root / "book.epub"
        write_epub(source)
        process, report, output = self.inspect(source, "epub")
        self.assertEqual(process.returncode, 0)
        self.assertEqual(report["status"], "partial")
        self.assertEqual(report["metadata"]["title"], "测试 EPUB")
        self.assertEqual(report["coverage"]["source_units"], 2)
        codes = {item["code"] for item in report["warnings"]}
        self.assertIn("empty_spine_items", codes)
        content = (output / "content.md").read_text(encoding="utf-8")
        self.assertLess(content.index("第一章"), content.index("第二章"))
        self.assertIn("OEBPS/chapter1.xhtml", content)

    def test_epub_non_body_roles_do_not_reduce_body_readability(self) -> None:
        source = self.root / "roles.epub"
        write_role_epub(source)
        process, report, output = self.inspect(source, "roles")
        self.assertEqual(process.returncode, 0)
        self.assertEqual(report["schema_version"], 2)
        self.assertEqual(report["status"], "readable")
        self.assertEqual(report["body_status"], "readable")
        self.assertEqual(report["coverage"]["role_counts"], {"cover": 1, "body": 1, "legal": 1})
        self.assertEqual(report["coverage"]["body_units"], 1)
        self.assertEqual(report["coverage"]["empty_body_units"], 0)
        self.assertLess(report["coverage"]["body_characters"], report["coverage"]["total_characters"])
        codes = {item["code"] for item in report["warnings"]}
        self.assertIn("empty_non_body_spine_items", codes)
        self.assertNotIn("empty_spine_items", codes)
        toc = (output / "toc.md").read_text(encoding="utf-8")
        self.assertIn("[cover] Cover", toc)
        self.assertIn("[body] 第一章", toc)
        self.assertIn("[legal] Project License", toc)

    def test_epub_suspicious_compression_is_rejected(self) -> None:
        source = self.root / "compressed.epub"
        write_role_epub(source)
        with zipfile.ZipFile(source, "a", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
            archive.writestr("OEBPS/suspicious.bin", b"0" * (8 * 1024 * 1024))
        process, report, output = self.inspect(source, "compressed")
        self.assertEqual(process.returncode, 2)
        self.assertEqual(report["status"], "unreadable")
        self.assertIn("epub_suspicious_compression", {item["code"] for item in report["warnings"]})
        self.assertFalse((output / "content.md").exists())

    def test_epub_missing_cover_does_not_reduce_body_readability(self) -> None:
        source = self.root / "missing-cover.epub"
        write_role_epub(source, include_cover=False)
        process, report, _ = self.inspect(source, "missing-cover")
        self.assertEqual(process.returncode, 0)
        self.assertEqual(report["status"], "readable")
        self.assertIn("unreadable_non_body_spine_items", {item["code"] for item in report["warnings"]})

    def test_damaged_epub_has_report_but_no_content(self) -> None:
        source = self.root / "broken.epub"
        source.write_bytes(b"not an epub container")
        process, report, output = self.inspect(source, "broken-epub")
        self.assertEqual(process.returncode, 2)
        self.assertEqual(report["status"], "unreadable")
        self.assertIn("damaged_epub", {item["code"] for item in report["warnings"]})
        self.assertFalse((output / "content.md").exists())

    def test_docx_headings_and_no_page_claim(self) -> None:
        from docx import Document

        source = self.root / "book.docx"
        document = Document()
        document.core_properties.title = "DOCX 书"
        document.core_properties.author = "测试作者"
        document.add_heading("第一章", level=1)
        document.add_paragraph("这是第一章正文。")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "概念"
        table.rows[0].cells[1].text = "解释"
        document.save(source)

        _, report, output = self.inspect(source, "docx")
        self.assertEqual(report["metadata"]["title"], "DOCX 书")
        self.assertEqual(report["directory_status"], "available")
        content = (output / "content.md").read_text(encoding="utf-8")
        self.assertIn("DOCX: 第一章", content)
        self.assertNotIn("DOCX 第 1 页", content)

    def make_pdf(self, path: Path, text: str | None) -> None:
        from reportlab.pdfgen import canvas

        pdf = canvas.Canvas(str(path))
        if text:
            pdf.drawString(72, 720, text)
        pdf.showPage()
        pdf.save()

    def test_pdf_uses_actual_page_locator(self) -> None:
        source = self.root / "book.pdf"
        self.make_pdf(source, "Readable PDF book content for page locator testing.")
        _, report, output = self.inspect(source, "pdf")
        self.assertEqual(report["format"], "pdf")
        self.assertEqual(report["locator_scheme"], "actual PDF page number")
        self.assertIn("PDF 第 1 页", (output / "content.md").read_text(encoding="utf-8"))

    def test_blank_pdf_is_flagged_as_possible_scan(self) -> None:
        source = self.root / "scan.pdf"
        self.make_pdf(source, None)
        process, report, output = self.inspect(source, "scan")
        self.assertEqual(process.returncode, 2)
        self.assertEqual(report["status"], "unreadable")
        self.assertIn("possible_scanned_pdf", {item["code"] for item in report["warnings"]})
        self.assertFalse((output / "content.md").exists())

    def test_encrypted_pdf_has_no_fake_content(self) -> None:
        from pypdf import PdfReader, PdfWriter

        plain = self.root / "plain.pdf"
        source = self.root / "encrypted.pdf"
        self.make_pdf(plain, "Secret book content")
        reader = PdfReader(str(plain))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt("secret")
        with source.open("wb") as handle:
            writer.write(handle)

        process, report, output = self.inspect(source, "encrypted")
        self.assertEqual(process.returncode, 2)
        self.assertEqual(report["status"], "unreadable")
        self.assertIn("encrypted_pdf", {item["code"] for item in report["warnings"]})
        self.assertFalse((output / "content.md").exists())

    def test_control_characters_are_reported(self) -> None:
        source = self.root / "garbled.txt"
        source.write_bytes("正常内容\x01仍可读取".encode("utf-8"))
        _, report, _ = self.inspect(source, "garbled")
        self.assertIn("control_characters", {item["code"] for item in report["warnings"]})

    def test_dependency_preflight_is_machine_readable(self) -> None:
        process = subprocess.run(
            [sys.executable, str(SCRIPT), "--check-dependencies"],
            text=True,
            capture_output=True,
            check=False,
        )
        self.assertEqual(process.returncode, 0, process.stderr)
        report = json.loads(process.stdout)
        self.assertIn(report["status"], {"ready", "partial"})
        self.assertIn("pypdf", {item["module"] for item in report["optional_dependencies"].values()})
        self.assertIn("docx", {item["module"] for item in report["optional_dependencies"].values()})


if __name__ == "__main__":
    unittest.main(verbosity=2)
